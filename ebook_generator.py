
import openai
import os
import time
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import re

# Enter your API key and Folder path
openai.api_key = "API-KEY"
user_path = "C:/Users/Name/Desktop/"

# Enter the Title of the Ebook
title = "Lose Weight"

doc = Document()

title_text = doc.add_paragraph()
run = title_text.add_run(title)
run.bold = True
run.font.size = Pt(25)
title_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Outline
try:
    outline_response = openai.Completion.create(
        engine="text-davinci-003",
        prompt= title+' is the ebook title, give me the outline of the ebook with introduction and additional 5 chapters and 3 detailed subchapters for each chapter',
        max_tokens=4010,
        n=1,
        stop=None,
        temperature=1.0,
    )

    outline_text = doc.add_paragraph()
    outline = outline_response['choices'][0]['text']
    run = outline_text.add_run(outline)
    run.font.size = Pt(13)

    doc.add_page_break()

    folder = user_path+title+'/'
    if not os.path.exists(folder):
        os.makedirs(folder)
    ebook = folder+title+" Ebook.docx"
    doc.save(ebook)

    time.sleep(20)

except Exception as e:
    print("Error Occurred: ", e)

# Disclaimer
try:
    disclaimer_response = openai.Completion.create(
        engine="text-davinci-003",
        prompt= title+' is the ebook title, write a Disclaimer for the ebook',
        max_tokens=2010,
        n=1,
        stop=None,
        temperature=1.0,
    )

    disclaimer_text = doc.add_paragraph()
    run = disclaimer_text.add_run("Disclaimer")
    run.bold = True
    run.font.size = Pt(18)

    disclaimer_content = doc.add_paragraph()
    disclaimer = disclaimer_response['choices'][0]['text']
    run = disclaimer_content.add_run(disclaimer)
    run.font.size = Pt(13)

    doc.add_page_break()

    doc.save(ebook)

    time.sleep(20)
    
except Exception as e:
    print("Error Occurred 1: ", e)


section = doc.sections[0]
section.page_height = Inches(10)
section.page_width = Inches(7)

doc.styles['Normal'].font.name = 'Lato'

section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)


chapters_and_subchapters = re.findall(r'Chapter \d+: (.+?)\n((?:.*?\n){3})', outline)

# (chapter_title, [subchapter_lines])
result = [(title.strip(), subchapters.strip().split('\n')) for title, subchapters in chapters_and_subchapters]
chapters, subchapters_list = zip(*result)

subchapters = [[re.sub(r'^[\d.\s-]+', '', subchapter) for subchapter in chapter] for chapter in subchapters_list]


for i,chapter in enumerate(chapters):

    # Add Chapter
    paragraph1 = doc.add_paragraph()
    run1 = paragraph1.add_run(chapter)
    run1.bold = True
    run1.font.size = Pt(22)
    paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    for subchapter in subchapters[i]:
        
        # Subchapter 
        subchapter_text = doc.add_paragraph()
        run = subchapter_text.add_run(subchapter)
        run.bold = True
        run.font.size = Pt(18)

        response1 = openai.Completion.create(
            engine="text-davinci-003",
            prompt="Write the full chapter of "+subchapter+" of the ebook "+title+" in 1000 words",
            max_tokens=4048,
            n=1,
            stop=None,
            temperature=1.0,
        )

        # Add Subchapter Content
        subchapter_content1 = doc.add_paragraph()
        subchapter_1 = response1['choices'][0]['text']
        run = subchapter_content1.add_run(subchapter_1)
        run.font.size = Pt(13)

        response2 = openai.Completion.create(
            engine="text-davinci-003",
            prompt="Continue the chapter "+subchapter+" of the ebook "+title,
            max_tokens=4048,
            n=1,
            stop=None,
            temperature=1.0,
        )

        # Add Extra Subchapter Content
        subchapter_content2 = doc.add_paragraph()
        subchapter_2 = response2['choices'][0]['text']
        run = subchapter_content2.add_run(subchapter_2)
        run.font.size = Pt(13)

        doc.add_page_break()

        time.sleep(25)

    n=n+1
    
    doc.save(ebook)

# Conclusion
conclusion_text = doc.add_paragraph()
run = conclusion_text.add_run("Conclusion")
run.bold = True
run.font.size = Pt(18)
conclusion_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

conclusion_response = openai.Completion.create(
    engine="text-davinci-003",
    prompt= title+' is the ebook title, write a detailed Conclusion for the ebook',
    max_tokens=4000,
    n=1,
    stop=None,
    temperature=1.0,
)

conclusion_content1 = doc.add_paragraph()
conclusion_1 = conclusion_response['choices'][0]['text']
run = conclusion_content1.add_run(conclusion_1)
run.font.size = Pt(13)

# Prompt for the conclusion continution of the ebook
conclusion_response2 = openai.Completion.create(
    engine="text-davinci-003",
    prompt= 'Continue the Conclusion for the ebook '+title,
    max_tokens=4000,
    n=1,
    stop=None,
    temperature=1.0,
)

conclusion_content2 = doc.add_paragraph()
conclusion_2 = conclusion_response2['choices'][0]['text']
run = conclusion_content2.add_run(conclusion_2)
run.font.size = Pt(13)

doc.save(ebook)
os.startfile(ebook)